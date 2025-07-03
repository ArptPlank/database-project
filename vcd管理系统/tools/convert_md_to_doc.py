#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换工具
支持批量转换markdown文件为docx格式
"""

import os
import sys
import subprocess
from pathlib import Path
import argparse

def check_pandoc():
    """检查pandoc是否安装"""
    try:
        result = subprocess.run(['pandoc', '--version'], 
                              capture_output=True, text=True, check=True)
        print("✅ Pandoc已安装:", result.stdout.split('\n')[0])
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("❌ 未找到Pandoc，请先安装Pandoc")
        print("安装方法:")
        print("Windows: 下载安装包 https://pandoc.org/installing.html")
        print("macOS: brew install pandoc")
        print("Linux: sudo apt-get install pandoc")
        return False

def convert_with_pandoc(md_file, output_file=None):
    """使用pandoc转换markdown到docx"""
    if not output_file:
        output_file = Path(md_file).with_suffix('.docx')
    
    try:
        # 使用pandoc转换，支持中文
        cmd = [
            'pandoc',
            str(md_file),
            '-o', str(output_file),
            '--from=markdown',
            '--to=docx',
            '--standalone',
            '--toc',  # 生成目录
            '--reference-doc=template.docx'  # 如果有模板的话
        ]
        
        # 如果没有模板文件，去掉这个参数
        if not Path('template.docx').exists():
            cmd = cmd[:-1]
        
        subprocess.run(cmd, check=True)
        print(f"✅ 转换成功: {md_file} -> {output_file}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"❌ 转换失败: {md_file}")
        print(f"错误信息: {e}")
        return False

def convert_with_python_libs(md_file, output_file=None):
    """使用Python库转换markdown到docx"""
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
    except ImportError as e:
        print(f"❌ 缺少必要的Python库: {e}")
        print("请安装: pip install markdown python-docx")
        return False
    
    if not output_file:
        output_file = Path(md_file).with_suffix('.docx')
    
    try:
        # 读取markdown文件
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        setup_document_styles(doc)
        
        # 解析并转换markdown内容
        parse_markdown_to_docx(md_content, doc)
        
        # 保存文档
        doc.save(output_file)
        print(f"✅ 转换成功: {md_file} -> {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ 转换失败: {md_file}")
        print(f"错误信息: {e}")
        return False

def setup_document_styles(doc):
    """设置Word文档样式"""
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 创建标题样式
    for i in range(1, 7):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = '微软雅黑'
            heading_font.size = Pt(18 - i * 2)
        except:
            pass

def parse_markdown_to_docx(md_content, doc):
    """解析markdown内容并添加到Word文档"""
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('# ').strip()
            if level <= 6:
                heading = doc.add_heading(title, level=level)
            else:
                doc.add_paragraph(title)
                
        # 处理代码块
        elif line.startswith('```'):
            continue  # 简化处理，跳过代码块标记
            
        # 处理列表
        elif line.startswith('-') or line.startswith('*'):
            text = line.lstrip('- *').strip()
            doc.add_paragraph(text, style='List Bullet')
            
        # 处理编号列表
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')
            
        # 处理普通段落
        else:
            if line:
                doc.add_paragraph(line)

def merge_all_reports():
    """合并所有报告文件为一个完整的Word文档"""
    report_files = [
        "数据库课程设计报告.md",
        "第三四章-概念逻辑设计.md", 
        "第七章-应用系统设计与实现.md",
        "第九十章-测试与总结.md",
        "附录-完整报告.md"
    ]
    
    # 检查文件是否存在
    existing_files = []
    for file in report_files:
        if Path(file).exists():
            existing_files.append(file)
        else:
            print(f"⚠️  文件不存在: {file}")
    
    if not existing_files:
        print("❌ 没有找到任何报告文件")
        return False
    
    print(f"📄 找到 {len(existing_files)} 个报告文件")
    
    # 方法1: 使用pandoc合并转换
    if check_pandoc():
        try:
            cmd = [
                'pandoc'
            ] + existing_files + [
                '-o', '完整数据库课程设计报告.docx',
                '--from=markdown',
                '--to=docx',
                '--standalone',
                '--toc',
                '--toc-depth=3'
            ]
            
            subprocess.run(cmd, check=True)
            print("✅ 完整报告生成成功: 完整数据库课程设计报告.docx")
            return True
        except subprocess.CalledProcessError as e:
            print(f"❌ Pandoc转换失败: {e}")
    
    # 方法2: 使用Python库合并
    try:
        from docx import Document
        
        doc = Document()
        setup_document_styles(doc)
        
        # 添加标题页
        title = doc.add_heading('音像店VCD管理系统数据库课程设计报告', 0)
        doc.add_paragraph('')
        doc.add_paragraph('课程名称：数据库原理与应用')
        doc.add_paragraph('设计题目：音像店VCD管理系统')
        doc.add_paragraph('技术栈：MongoDB + Python Flask + HTML/CSS/JavaScript')
        doc.add_paragraph(f'完成时间：2025年7月')
        doc.add_page_break()
        
        # 合并所有文件内容
        for file_path in existing_files:
            print(f"📖 处理文件: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 添加文件分隔
            doc.add_heading(f"来源文件: {file_path}", level=1)
            parse_markdown_to_docx(content, doc)
            doc.add_page_break()
        
        doc.save('完整数据库课程设计报告.docx')
        print("✅ 完整报告生成成功: 完整数据库课程设计报告.docx")
        return True
        
    except ImportError:
        print("❌ 缺少python-docx库，请安装: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ Python库转换失败: {e}")
        return False

def convert_single_file(md_file, use_pandoc=True):
    """转换单个文件"""
    if not Path(md_file).exists():
        print(f"❌ 文件不存在: {md_file}")
        return False
    
    if use_pandoc and check_pandoc():
        return convert_with_pandoc(md_file)
    else:
        return convert_with_python_libs(md_file)

def main():
    parser = argparse.ArgumentParser(description='Markdown到Word文档转换工具')
    parser.add_argument('--file', '-f', help='指定要转换的markdown文件')
    parser.add_argument('--merge', '-m', action='store_true', help='合并所有报告文件')
    parser.add_argument('--use-python', '-p', action='store_true', help='使用Python库而不是pandoc')
    parser.add_argument('--all', '-a', action='store_true', help='转换所有markdown文件')
    
    args = parser.parse_args()
    
    print("🔄 Markdown到Word转换工具")
    print("=" * 50)
    
    if args.merge:
        print("📋 合并所有报告文件...")
        merge_all_reports()
    elif args.file:
        print(f"📄 转换单个文件: {args.file}")
        convert_single_file(args.file, not args.use_python)
    elif args.all:
        print("📚 转换所有markdown文件...")
        md_files = list(Path('.').glob('*.md'))
        if not md_files:
            print("❌ 当前目录没有找到markdown文件")
            return
        
        success_count = 0
        for md_file in md_files:
            if convert_single_file(md_file, not args.use_python):
                success_count += 1
        
        print(f"\n✅ 转换完成: {success_count}/{len(md_files)} 个文件转换成功")
    else:
        # 默认行为：合并所有报告文件
        print("📋 默认操作：合并所有报告文件...")
        print("如需其他操作，请使用 --help 查看帮助")
        merge_all_reports()

if __name__ == "__main__":
    main() 